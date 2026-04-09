import datetime
import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

FIGURE_BLOCK_PATTERN = r"<<FIGURE>>\s*.*?\s*<</FIGURE>>"
CITATION_PATTERN = re.compile(r"\[(\d{1,3})\]")


def _set_run_font(
    run,
    zh_font: str = "宋体",
    en_font: str = "Times New Roman",
    size_pt: float | None = None,
    bold: bool | None = None,
    underline: bool | None = None,
    color_rgb: RGBColor | None = None,
) -> None:
    """统一设置 run 的中英文字体、字号和样式。"""
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), zh_font)
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    run.font.name = en_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def _apply_fixed_line_spacing(paragraph_format, pt: float = 22) -> None:
    """应用固定值行距。"""
    p_pr = paragraph_format._element.get_or_add_pPr()
    for old in p_pr.findall(qn("w:spacing")):
        p_pr.remove(old)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)


def _restart_page_numbering(section, start: int = 1) -> None:
    """为当前 section 重置页码起始值。"""
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(old)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start))
    sect_pr.append(pg_num)


# ---------- TOC: heading pre-scan / bookmark / PAGEREF ----------


def _pre_scan_headings(
    full_text: str,
    title: str = "",
) -> list[dict[str, object]]:
    """Pre-scan body Markdown to extract level 1-3 headings for the TOC page.

    Returns a list like::

        [
            {"text": "First Chapter", "level": 1, "bookmark": "_toc_0", "bookmark_id": 100},
            {"text": "1.1 Background", "level": 2, "bookmark": "_toc_1", "bookmark_id": 101},
        ]
    """
    clean_text = re.sub(FIGURE_BLOCK_PATTERN, "", full_text, flags=re.DOTALL)
    entries: list[dict[str, object]] = []

    # Non-body headings that should never appear in the TOC,
    # even if the LLM accidentally generates them inside the body text.
    _NON_BODY_HEADINGS = {
        "摘要", "摘 要", "中文摘要",
        "abstract",
        "致谢", "致 谢",
        "参考文献",
    }

    for line in clean_text.split("\n"):
        line = line.strip()
        if not line or line == "---pagebreak---":
            continue

        level = 0
        text = ""
        if line.startswith("### "):
            level, text = 3, line[4:]
        elif line.startswith("## "):
            level, text = 2, line[3:]
        elif line.startswith("# "):
            level, text = 1, line[2:]

        if level == 0:
            continue

        text = text.strip()
        # Exclude the thesis title itself
        if title and text == title.strip():
            continue
        # Exclude non-body section headings (defensive against prompt drift)
        if text.lower() in _NON_BODY_HEADINGS:
            continue

        idx = len(entries)
        entries.append(
            {
                "text": text,
                "level": level,
                "bookmark": f"_toc_{idx}",
                "bookmark_id": idx + 100,
            }
        )

    return entries


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    """Wrap paragraph content with bookmarkStart / bookmarkEnd."""
    p_element = paragraph._element

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))

    p_pr = p_element.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.addnext(bm_start)
    else:
        p_element.insert(0, bm_start)

    p_element.append(bm_end)


def _add_pageref_field(paragraph, bookmark_name: str, cached_page: str = "?") -> None:
    """Append a PAGEREF field referencing *bookmark_name* to *paragraph*.

    *cached_page* is the pre-estimated page number displayed until the
    end-user's application recalculates fields.  ``dirty="true"`` on the
    begin fldChar signals Word / WPS that this field should be refreshed.

    The generated OOXML structure::

        <w:fldChar begin dirty="true"/>
        <w:instrText> PAGEREF _toc_0 \\h </w:instrText>
        <w:fldChar separate/>
        <w:t>5</w:t>          (estimated; corrected on F9 / auto-update)
        <w:fldChar end/>
    """
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run_begin._element.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    run_instr._element.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._element.append(fld_sep)

    run_placeholder = paragraph.add_run(str(cached_page))
    _set_run_font(run_placeholder, size_pt=12)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._element.append(fld_end)




def _clear_header_footer(section) -> None:
    """清空当前 section 的页眉页脚。"""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        paragraph.clear()
    for paragraph in section.footer.paragraphs:
        paragraph.clear()


def _copy_page_layout_from_first(document: Document, section) -> None:
    """让新增 section 继承第一页的纸张和页边距设置。"""
    first = document.sections[0]
    section.page_width = first.page_width
    section.page_height = first.page_height
    section.top_margin = first.top_margin
    section.bottom_margin = first.bottom_margin
    section.left_margin = first.left_margin
    section.right_margin = first.right_margin


def _make_blank_section(document: Document):
    """新增一个空白 section，不带页眉页脚。"""
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _copy_page_layout_from_first(document, section)
    _clear_header_footer(section)
    return section


def _setup_body_section(document: Document, title: str) -> None:
    """配置正文 section 的页眉和页脚页码。"""
    section = document.sections[-1]
    _copy_page_layout_from_first(document, section)
    _clear_header_footer(section)
    _restart_page_numbering(section, start=1)

    header = section.header
    p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_header.clear()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run(title)
    _set_run_font(run_h, size_pt=10)

    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.clear()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_f = p_footer.add_run()
    r_element = run_f._element

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r_element.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_element.append(fld_sep)

    num_run = p_footer.add_run("1")
    _set_run_font(
        num_run,
        zh_font="Times New Roman",
        en_font="Times New Roman",
        size_pt=10,
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p_footer.add_run()._element.append(fld_end)


# ---------- Markdown 表格解析工具 ----------


def _is_table_separator(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行。"""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.match(r"^[-:]+$", c) for c in cells) if cells else False


def _parse_table_line(line: str) -> list[str]:
    """解析单行 Markdown 表格。"""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _collect_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """收集连续的 Markdown 表格行。"""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        if _is_table_separator(lines[i]):
            i += 1
            continue
        rows.append(_parse_table_line(lines[i]))
        i += 1
    return rows, i


def _add_markdown_text_to_paragraph(
    paragraph,
    text: str,
    is_header: bool = False,
    is_table: bool = False,
) -> None:
    """解析 Markdown 内联加粗并写入段落。"""
    text = text.replace("~", "")
    parts = re.split(r"(\*\*.*?\*\*|\[\d{1,3}\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif CITATION_PATTERN.fullmatch(part):
            run = paragraph.add_run(part)
            run.font.superscript = True
        else:
            run = paragraph.add_run(part)
            if is_header:
                run.bold = True
        _set_run_font(run, zh_font="宋体")
        if is_table:
            run.font.size = Pt(10.5)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    """将二维数据写入 Word 表格。"""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=num_cols, style="Table Grid")

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                continue
            cell = table.rows[i].cells[j]
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            _add_markdown_text_to_paragraph(
                paragraph,
                cell_text,
                is_header=(i == 0),
                is_table=True,
            )


# ---------- 文档样式初始化 ----------


def _init_styles(document: Document) -> None:
    """设置文档默认样式和标题样式。"""
    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.first_line_indent = Cm(0.74)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(6)
    _apply_fixed_line_spacing(normal_style.paragraph_format, pt=22)

    normal_r_pr = normal_style.element.get_or_add_rPr()
    normal_r_fonts = normal_r_pr.get_or_add_rFonts()
    normal_r_fonts.set(qn("w:eastAsia"), "宋体")
    normal_r_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_r_fonts.set(qn("w:hAnsi"), "Times New Roman")

    heading_config = [
        (1, 18, True),
        (2, 16, True),
        (3, 14, True),
    ]
    for level, size, bold in heading_config:
        style = document.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), "黑体"
        )
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------- 页面设置 ----------


def _setup_page(document: Document) -> None:
    """设置首页 section 的页面尺寸和页边距。"""
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    _clear_header_footer(section)


# ---------- 固定页面 ----------


def _add_cover_page(
    document: Document,
    title: str,
    author: str,
    advisor: str,
    degree_type: str,
    major: str,
    school: str,
    year_month: str,
) -> None:
    """构建封面页。"""

    def _blank(lines: int) -> None:
        for _ in range(lines):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    _blank(6)

    p_main = document.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main.paragraph_format.first_line_indent = Pt(0)
    # Cover title uses a much larger font than body text; if it inherits the
    # Normal style's fixed 22pt line height, WPS clips the glyphs vertically.
    p_main.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_main = p_main.add_run("学士学位论文")
    _set_run_font(run_main, zh_font="黑体", size_pt=36, bold=True)
    p_main.paragraph_format.space_after = Pt(24)

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_title = p_title.add_run(title)
    _set_run_font(run_title, size_pt=18, underline=True)
    p_title.paragraph_format.space_after = Pt(36)

    fields = [
        ("作者姓名", author),
        ("指导教师", advisor),
        ("学位类别", degree_type),
        ("专    业", major),
        ("学院（系）", school),
    ]
    for label, value in fields:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(3.5)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run_label = paragraph.add_run(f"{label}：")
        _set_run_font(run_label, size_pt=14)
        run_value = paragraph.add_run(value)
        _set_run_font(run_value, size_pt=14, bold=True)

    _blank(8)

    if not year_month:
        now = datetime.datetime.now()
        year_month = f"{now.year} 年 {now.month} 月"

    p_date = document.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.first_line_indent = Pt(0)
    run_date = p_date.add_run(year_month)
    _set_run_font(run_date, size_pt=14, bold=True)

    document.add_page_break()


def _add_abstract_zh_page(document: Document, abstract: str, keywords: str) -> None:
    """中文摘要页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("摘  要")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    if abstract:
        for para_text in abstract.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(24)
            run = paragraph.add_run(para_text)
            _set_run_font(run, size_pt=12)

    if keywords:
        p_kw = document.add_paragraph()
        p_kw.paragraph_format.first_line_indent = Pt(0)
        p_kw.paragraph_format.space_before = Pt(12)
        run_label = p_kw.add_run("关键词：")
        _set_run_font(run_label, size_pt=12, bold=True)
        run_kw = p_kw.add_run(keywords)
        _set_run_font(run_kw, size_pt=12)

    document.add_page_break()


def _add_abstract_en_page(document: Document, abstract: str, keywords: str) -> None:
    """英文摘要页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("Abstract")
    _set_run_font(
        run_title,
        zh_font="Times New Roman",
        en_font="Times New Roman",
        size_pt=16,
        bold=True,
    )

    if abstract:
        for para_text in abstract.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            run = paragraph.add_run(para_text)
            _set_run_font(
                run,
                zh_font="Times New Roman",
                en_font="Times New Roman",
                size_pt=12,
            )

    if keywords:
        p_kw = document.add_paragraph()
        p_kw.paragraph_format.first_line_indent = Pt(0)
        p_kw.paragraph_format.space_before = Pt(12)
        run_label = p_kw.add_run("Keywords: ")
        _set_run_font(
            run_label,
            zh_font="Times New Roman",
            en_font="Times New Roman",
            size_pt=12,
            bold=True,
        )
        run_kw = p_kw.add_run(keywords)
        _set_run_font(
            run_kw,
            zh_font="Times New Roman",
            en_font="Times New Roman",
            size_pt=12,
        )


def _estimate_page_numbers(
    full_text: str,
    toc_entries: list[dict[str, object]],
    body_start_page: int = 1,
) -> dict[str, int]:
    """Estimate the page number for each TOC heading.

    Uses a simple character-count heuristic:
      - A4 with 3cm/2.5cm margins, 12pt Song, 22pt exact line spacing
      - ~32 lines × ~35 chars ≈ **1100 chars per page**
      - Explicit ``---pagebreak---`` forces a new page
      - Each level-1 heading (``# ``) is preceded by a page break in the
        document (except the first), so it always starts a new page.

    Returns ``{bookmark_name: estimated_page, ...}``.
    """
    CHARS_PER_PAGE = 1100
    clean = re.sub(FIGURE_BLOCK_PATTERN, "", full_text, flags=re.DOTALL)

    page = body_start_page
    char_count = 0
    page_map: dict[str, int] = {}

    # Build a quick lookup: heading_text -> bookmark (first match wins)
    heading_queue: list[dict[str, object]] = list(toc_entries)
    hq_idx = 0

    for line in clean.split("\n"):
        stripped = line.strip()

        # Explicit page break
        if stripped == "---pagebreak---":
            page += 1
            char_count = 0
            continue

        # Detect heading level
        level = 0
        text = ""
        if stripped.startswith("### "):
            level, text = 3, stripped[4:].strip()
        elif stripped.startswith("## "):
            level, text = 2, stripped[3:].strip()
        elif stripped.startswith("# "):
            level, text = 1, stripped[2:].strip()

        # Match heading to next expected TOC entry
        if level > 0 and hq_idx < len(heading_queue):
            entry = heading_queue[hq_idx]
            if entry["text"] == text and entry["level"] == level:
                page_map[str(entry["bookmark"])] = page
                hq_idx += 1

        # Accumulate characters for page estimation
        char_count += len(stripped) + 1  # +1 for newline equivalent
        if char_count >= CHARS_PER_PAGE:
            page += 1
            char_count = 0

    return page_map


def _add_toc_page(
    document: Document,
    toc_entries: list[dict[str, object]],
    full_text: str = "",
) -> None:
    """Generate a visible TOC page with PAGEREF dynamic page numbers.

    Each entry is a normal paragraph structured as::

        Heading text .................. page_number

    Page numbers use PAGEREF fields referencing bookmarks on body headings,
    so WPS/Word calculates the correct page on open.  Pre-estimated page
    numbers are written as cached values so the TOC is immediately readable
    even when the application does not auto-update fields.
    """
    # Estimate how many TOC pages we'll need (affects body start page)
    toc_page_count = max(1, (len(toc_entries) + 16) // 17)  # ~17 entries/page
    # Front matter: cover(1) + abstract_zh(1) + abstract_en(1) + TOC pages
    body_start_page = 1  # body section restarts page numbering from 1
    page_map = _estimate_page_numbers(full_text, toc_entries, body_start_page)

    # ---- TOC title ----
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("目  录")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    # Per-level left indentation
    indent_map = {1: Cm(0), 2: Cm(0.74), 3: Cm(1.48)}

    for entry in toc_entries:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = indent_map.get(int(entry["level"]), Cm(0))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _apply_fixed_line_spacing(p.paragraph_format, pt=22)

        # Right-aligned tab stop with dot leader at right margin (15.5 cm)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )

        # Heading text
        run_text = p.add_run(str(entry["text"]))
        _set_run_font(run_text, size_pt=12)
        if entry["level"] == 1:
            run_text.bold = True

        # Tab character (triggers dot leader extending to page number)
        tab_run = p.add_run("\t")
        _set_run_font(tab_run, size_pt=12)

        # PAGEREF field with estimated page number as cached value
        bm = str(entry["bookmark"])
        estimated = str(page_map.get(bm, "1"))
        _add_pageref_field(p, bm, cached_page=estimated)

    # Document setting: auto-update all fields on open (including PAGEREF)
    doc_settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc_settings.append(update_fields)



def _add_acknowledgment_page(document: Document, acknowledgment: str) -> None:
    """致谢页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("致  谢")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    if not acknowledgment:
        return

    for para_text in acknowledgment.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(24)
        run = paragraph.add_run(para_text)
        _set_run_font(run, size_pt=12)


def _add_references_page(document: Document, references: str) -> None:
    """参考文献页。"""
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("参考文献")
    _set_run_font(run_title, zh_font="黑体", size_pt=16, bold=True)

    if not references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run("（参考文献未生成，可能因未配置检索服务或检索结果不足）")
        _set_run_font(run, size_pt=12)
        return

    for line in references.splitlines():
        line = line.strip()
        if not line:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(line)
        _set_run_font(run, size_pt=12)


def _insert_picture_with_constraints(
    document: Document,
    image_path: str,
    max_width_cm: float = 15.5,
    max_height_cm: float = 9.5,
) -> None:
    """
    插入图片并限制最大宽高，避免图片过高导致版面被大面积占用。
    保持原始纵横比，不做拉伸。
    """
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    if width_px <= 0 or height_px <= 0:
        document.add_picture(image_path, width=Cm(max_width_cm))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    width_cm = max_width_cm
    height_cm = width_cm * height_px / width_px
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width_px / height_px

    document.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 强制单倍行距，避免被 Normal 样式的固定值（22pt）截断导致图片被文本覆盖
    from docx.enum.text import WD_LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Pt(0)


# ---------- 主构建函数 ----------


def build_word_document(
    full_text: str,
    placeholders: list[dict],
    image_paths: dict[int, str | None],
    output_path: str = "app/output/thesis.docx",
    title: str = "论文题目",
    author: str = "作者姓名",
    advisor: str = "指导教师",
    degree_type: str = "学士",
    major: str = "专业名称",
    school: str = "XX大学XX学院",
    year_month: str = "",
    abstract_zh: str = "",
    abstract_en: str = "",
    keywords_zh: str = "",
    keywords_en: str = "",
    acknowledgment: str = "",
    references: str = "",
) -> str:
    """将论文正文、图片与前后置页面构造成 Word 文档。"""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    _setup_page(document)
    _init_styles(document)

    # Section 1：封面 + 中英文摘要
    _add_cover_page(
        document,
        title,
        author,
        advisor,
        degree_type,
        major,
        school,
        year_month,
    )
    _add_abstract_zh_page(document, abstract_zh, keywords_zh)
    _add_abstract_en_page(document, abstract_en, keywords_en)

    # 预扫描正文标题，生成目录条目列表
    toc_entries = _pre_scan_headings(full_text, title=title)

    # Section 2：目录（可见条目 + PAGEREF 动态页码）
    _make_blank_section(document)
    _add_toc_page(document, toc_entries, full_text=full_text)

    # Section 3：正文 + 致谢 + 参考文献
    _make_blank_section(document)
    _setup_body_section(document, title)

    segments = re.split(FIGURE_BLOCK_PATTERN, full_text, flags=re.DOTALL)
    placeholder_idx = 0

    # 目录条目迭代器，用于在正文标题上打 bookmark
    _toc_iter = iter(toc_entries)
    _next_toc = next(_toc_iter, None)

    for segment in segments:
        lines = [line.strip() for line in segment.strip().split("\n")]
        i = 0
        while i < len(lines):
            line = lines[i]

            if not line:
                i += 1
                continue

            if line == "---pagebreak---":
                document.add_page_break()
                i += 1
                continue

            if line.startswith("|") and "|" in line[1:]:
                rows, next_i = _collect_table_lines(lines, i)
                if rows:
                    _add_table(document, rows)
                i = next_i
                continue

            if re.match(r"^表\d+-\d+\s", line):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _add_markdown_text_to_paragraph(paragraph, line)
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=10)
                i += 1
                continue

            if line.startswith("### "):
                heading_text = line[4:]
                paragraph = document.add_heading(heading_text, level=3)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if (
                    _next_toc
                    and _next_toc["text"] == heading_text.strip()
                    and _next_toc["level"] == 3
                ):
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("## "):
                heading_text = line[3:]
                paragraph = document.add_heading(heading_text, level=2)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if (
                    _next_toc
                    and _next_toc["text"] == heading_text.strip()
                    and _next_toc["level"] == 2
                ):
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("# "):
                heading_text = line[2:]
                paragraph = document.add_heading(heading_text, level=1)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if (
                    _next_toc
                    and _next_toc["text"] == heading_text.strip()
                    and _next_toc["level"] == 1
                ):
                    _add_bookmark(paragraph, str(_next_toc["bookmark"]), int(_next_toc["bookmark_id"]))
                    _next_toc = next(_toc_iter, None)
            elif line.startswith("```"):
                pass
            elif line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _add_markdown_text_to_paragraph(paragraph, line[2:])
            elif re.match(r"^\d+\.\s+", line):
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.first_line_indent = Pt(0)
                text = re.sub(r"^\d+\.\s+", "", line)
                _add_markdown_text_to_paragraph(paragraph, text)
            else:
                paragraph = document.add_paragraph()
                _add_markdown_text_to_paragraph(paragraph, line)

            i += 1

        if placeholder_idx >= len(placeholders):
            continue

        placeholder = placeholders[placeholder_idx]
        image_path = image_paths.get(placeholder["index"])
        has_image = bool(image_path and Path(image_path).exists())

        if has_image and image_path:
            _insert_picture_with_constraints(document, image_path)

        caption = placeholder.get("caption", f"图{placeholder_idx + 1}")
        if not has_image:
            caption += "（图片生成失败）"
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.first_line_indent = Pt(0)
        caption_run = caption_paragraph.add_run(caption)
        _set_run_font(caption_run, size_pt=10)

        placeholder_idx += 1

    document.add_page_break()
    _add_acknowledgment_page(document, acknowledgment)

    document.add_page_break()
    _add_references_page(document, references)

    document.save(output_path)
    return output_path
